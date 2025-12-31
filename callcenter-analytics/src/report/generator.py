"""
Report Generator Module
Generate comprehensive PDF/DOCX reports
"""

import os
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd


class ReportGenerator:
    """Generate comprehensive analysis reports"""

    def __init__(self, config: dict):
        self.config = config
        self.report_config = config.get('report', {})
        self.output_path = self.report_config.get('output_path', 'outputs')
        os.makedirs(self.output_path, exist_ok=True)

    def create_docx_report(self, title: str, sections: Dict[str, any],
                          charts: Optional[Dict[str, str]] = None) -> str:
        """
        Create DOCX report

        Args:
            title: Report title
            sections: Dictionary of section_name: content
            charts: Optional dictionary of chart paths

        Returns:
            Path to generated report
        """
        print("\n📝 Generating DOCX report...")

        doc = Document()

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("=" * 80)

        # Add sections
        for section_name, content in sections.items():
            doc.add_heading(section_name, level=1)

            if isinstance(content, str):
                doc.add_paragraph(content)
            elif isinstance(content, dict):
                for key, value in content.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value))
            elif isinstance(content, list):
                for item in content:
                    doc.add_paragraph(str(item), style='List Bullet')

            doc.add_paragraph()  # Spacing

        # Add charts
        if charts:
            doc.add_page_break()
            doc.add_heading('Visualizations', level=1)

            for chart_name, chart_path in charts.items():
                if os.path.exists(chart_path):
                    doc.add_heading(chart_name.replace('_', ' ').title(), level=2)
                    doc.add_picture(chart_path, width=Inches(6))
                    doc.add_paragraph()

        # Save report
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"kpi_report_{timestamp}.docx"
        filepath = os.path.join(self.output_path, filename)
        doc.save(filepath)

        print(f"✓ Report saved: {filepath}")
        return filepath

    def create_text_report(self, title: str, sections: Dict[str, any]) -> str:
        """
        Create simple text report

        Args:
            title: Report title
            sections: Dictionary of section_name: content

        Returns:
            Path to generated report
        """
        print("\n📝 Generating text report...")

        lines = []
        lines.append("=" * 80)
        lines.append(title.center(80))
        lines.append("=" * 80)
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("=" * 80)
        lines.append("")

        for section_name, content in sections.items():
            lines.append(f"\n{'=' * 80}")
            lines.append(f"{section_name}")
            lines.append("=" * 80)

            if isinstance(content, str):
                lines.append(content)
            elif isinstance(content, dict):
                for key, value in content.items():
                    lines.append(f"  {key}: {value}")
            elif isinstance(content, list):
                for item in content:
                    lines.append(f"  • {item}")

            lines.append("")

        # Save report
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"kpi_report_{timestamp}.txt"
        filepath = os.path.join(self.output_path, filename)

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

        print(f"✓ Report saved: {filepath}")
        return filepath

    def create_comprehensive_report(self, analysis_data: Dict, insights: Dict,
                                   charts: Dict[str, str]) -> str:
        """
        Create comprehensive report with all analysis results

        Args:
            analysis_data: Analysis results (KPIs, correlations, etc.)
            insights: AI-generated insights
            charts: Dictionary of chart file paths

        Returns:
            Path to generated report
        """
        print("\n📝 Generating comprehensive report...")

        sections = {}

        # Executive Summary
        if 'executive_summary' in insights:
            sections['Executive Summary'] = insights['executive_summary']

        # KPI Overview
        if 'kpi_data' in analysis_data:
            kpi_data = analysis_data['kpi_data']
            kpi_text = self._format_kpi_data(kpi_data)
            sections['KPI Overview'] = kpi_text

        # Analysis Summary
        if 'summary' in insights:
            sections['Analysis Summary'] = insights['summary']

        # Pattern Analysis
        if 'patterns' in insights:
            sections['Identified Patterns'] = insights['patterns']

        # Recommendations
        if 'recommendations' in insights:
            sections['Recommendations'] = insights['recommendations']

        # Anomalies
        if 'anomalies' in insights:
            sections['Anomaly Analysis'] = insights['anomalies']

        # Statistical Summary
        if 'statistics' in analysis_data:
            sections['Statistical Summary'] = self._format_statistics(
                analysis_data['statistics']
            )

        # Generate report
        report_format = self.report_config.get('format', 'docx')

        if report_format == 'docx':
            filepath = self.create_docx_report(
                "Call Center KPI Analysis Report",
                sections,
                charts
            )
        else:
            filepath = self.create_text_report(
                "Call Center KPI Analysis Report",
                sections
            )

        print(f"\n✓ Comprehensive report generated: {filepath}")
        return filepath

    def _format_kpi_data(self, kpi_data: Dict) -> str:
        """Format KPI data as readable text"""
        lines = []

        for category, metrics in kpi_data.items():
            lines.append(f"\n{category.upper()}:")
            for metric, value in metrics.items():
                if isinstance(value, float):
                    lines.append(f"  {metric}: {value:.2f}")
                else:
                    lines.append(f"  {metric}: {value}")

        return '\n'.join(lines)

    def _format_statistics(self, stats: Dict) -> str:
        """Format statistical data as readable text"""
        lines = []

        if 'numeric' in stats and isinstance(stats['numeric'], pd.DataFrame):
            lines.append("\nNumeric Statistics:")
            lines.append(stats['numeric'].to_string())

        if 'categorical' in stats and isinstance(stats['categorical'], pd.DataFrame):
            lines.append("\n\nCategorical Statistics:")
            lines.append(stats['categorical'].to_string())

        return '\n'.join(lines)

    def create_summary_html(self, kpi_data: Dict, insights: Dict) -> str:
        """
        Create HTML summary for quick viewing

        Args:
            kpi_data: KPI data dictionary
            insights: Insights dictionary

        Returns:
            Path to HTML file
        """
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>KPI Analysis Summary</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #2c3e50; }}
                h2 {{ color: #34495e; margin-top: 30px; }}
                .metric {{ background: #ecf0f1; padding: 10px; margin: 5px 0; border-radius: 5px; }}
                .insight {{ background: #e8f5e9; padding: 15px; margin: 10px 0; border-left: 4px solid #4caf50; }}
            </style>
        </head>
        <body>
            <h1>Call Center KPI Analysis</h1>
            <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>

            <h2>Executive Summary</h2>
            <div class="insight">{insights.get('executive_summary', 'N/A')}</div>

            <h2>Key Metrics</h2>
        """

        for category, metrics in kpi_data.items():
            html_content += f"<h3>{category.title()}</h3>"
            for metric, value in list(metrics.items())[:5]:
                html_content += f'<div class="metric"><strong>{metric}:</strong> {value:.2f}</div>'

        html_content += """
        </body>
        </html>
        """

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"kpi_summary_{timestamp}.html"
        filepath = os.path.join(self.output_path, filename)

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)

        print(f"✓ HTML summary saved: {filepath}")
        return filepath
